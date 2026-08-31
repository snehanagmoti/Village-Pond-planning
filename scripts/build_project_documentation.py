"""Build the complete editable JalDrishti project documentation DOCX."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT = ROOT / "Final_Technical_Report.md"
ASSETS = ROOT / "tmp" / "docs_assets"
PRODUCTION = ROOT / "tmp" / "production" / "contour-auto.json"
OUTPUT = ROOT / "output" / "docx" / "JalDrishti_Complete_Project_Documentation.docx"

FONT = "Calibri"
MONO = "Consolas"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2B24"
GREEN = "0F7B63"
DEEP_GREEN = "0B5C4B"
MINT = "DDF8EE"
PALE_BLUE = "E9F6FB"
PALE_GOLD = "FFF7DE"
PALE_RED = "FDEAEA"
GRAY = "55756D"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name: str = FONT, size: float | None = None, *, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must sum to {TABLE_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def add_internal_link(paragraph, anchor: str, text: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_external_link(paragraph, url: str, text: str) -> None:
    relation_id = paragraph.part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


URL_PATTERN = re.compile(r"https?://[^\s>]+")
MD_LINK_PATTERN = re.compile(r"\[([^\]]+)\]\((https?://[^)]+)\)")


def clean_md(text: str) -> str:
    text = text.replace("**", "").replace("`", "")
    text = text.replace(" m2", " m²").replace(" m3", " m³")
    return text.strip()


def add_rich_text(paragraph, text: str, *, size: float = 11, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    cleaned = clean_md(text)
    cursor = 0
    tokens: list[tuple[int, int, str, str]] = []
    for match in MD_LINK_PATTERN.finditer(cleaned):
        tokens.append((match.start(), match.end(), match.group(1), match.group(2)))
    protected = [(start, end) for start, end, _, _ in tokens]
    for match in URL_PATTERN.finditer(cleaned):
        if any(start <= match.start() < end for start, end in protected):
            continue
        tokens.append((match.start(), match.end(), "Open link", match.group(0)))
    tokens.sort(key=lambda item: item[0])
    for start, end, label, url in tokens:
        if start > cursor:
            run = paragraph.add_run(cleaned[cursor:start])
            set_run_font(run, size=size, color=color, bold=bold, italic=italic)
        add_external_link(paragraph, url, label)
        cursor = end
    if cursor < len(cleaned):
        run = paragraph.add_run(cleaned[cursor:])
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    heading_tokens = {
        "Heading 1": (16, DEEP_GREEN, 18, 10),
        "Heading 2": (13, GREEN, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.page_break_before = False
    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(GRAY)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_with_next = False


def add_custom_numbering(doc: Document) -> tuple[int, int, int, int]:
    numbering = doc.part.numbering_part.element
    existing_abstract = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    next_abstract = max(existing_abstract or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def make_num(abstract_id: int, num_id: int, fmt: str, marker: str):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), marker)
        level.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        level.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        abstract.append(level)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        return abstract, num

    bullet_abstract, bullet_num = make_num(next_abstract, next_num, "bullet", "•")
    decimal_abstract, decimal_num = make_num(next_abstract + 1, next_num + 1, "decimal", "%1.")
    first_num_index = next((index for index, child in enumerate(numbering) if child.tag == qn("w:num")), len(numbering))
    numbering.insert(first_num_index, bullet_abstract)
    numbering.insert(first_num_index + 1, decimal_abstract)
    numbering.append(bullet_num)
    numbering.append(decimal_num)
    return next_num, next_num + 1, next_abstract, next_abstract + 1


def clone_numbering(doc: Document, abstract_id: int) -> int:
    numbering = doc.part.numbering_part.element
    existing_num = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_num or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)


def add_bullet(doc: Document, text: str, bullet_id: int) -> None:
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, bullet_id)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    add_rich_text(paragraph, text)


def add_number(doc: Document, text: str, number_id: int) -> None:
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, number_id)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    add_rich_text(paragraph, text)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[str]], widths: list[int] | None = None, *, font_size: float = 8.8) -> None:
    rows_list = [list(row) for row in rows]
    columns = len(headers)
    if widths is None:
        base = TABLE_WIDTH_DXA // columns
        widths = [base] * columns
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table = doc.add_table(rows=1, cols=columns)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    mark_header_row(header)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        set_cell_shading(cell, MINT)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(clean_md(value))
        set_run_font(run, size=font_size, color=INK, bold=True)
    for row_values in rows_list:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            add_rich_text(paragraph, str(value), size=font_size)
    citation = doc.add_paragraph()
    citation.paragraph_format.space_before = Pt(4)
    citation.paragraph_format.space_after = Pt(4)


def add_callout(doc: Document, title: str, text: str, *, fill: str = MINT, border: str = GREEN) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, size=10.5, color=border, bold=True)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    add_rich_text(paragraph, text, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    mark_header_row(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F6F8FA")
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(code.rstrip())
    set_run_font(run, MONO, 8.2, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_figure(doc: Document, path: Path, caption: str, alt_text: str, figure_number: int, *, width: float = 6.3) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    caption_paragraph = doc.add_paragraph(style="Caption")
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.add_run(f"Figure {figure_number}. {caption}")


def add_heading(doc: Document, text: str, level: int, bookmark_name: str, bookmark_id: int, *, page_break: bool = False) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    if page_break:
        paragraph.paragraph_format.page_break_before = True
    paragraph.add_run(clean_md(text))
    add_bookmark(paragraph, bookmark_name, bookmark_id)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("JalDrishti | Complete Project Documentation")
    set_run_font(run, size=8.5, color=GRAY, bold=True)
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6500, 2860])
    mark_header_row(table.rows[0])
    table.style = "Table Grid"
    for cell in table.row_cells(0):
        cell._tc.get_or_add_tcPr().remove(cell._tc.get_or_add_tcPr().find(qn("w:tcBorders"))) if cell._tc.get_or_add_tcPr().find(qn("w:tcBorders")) is not None else None
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("Screening decision support - field verification required")
    set_run_font(run, size=8, color=GRAY)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    run = right.add_run("Page ")
    set_run_font(run, size=8, color=GRAY)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    for field_node in (field_begin, instruction, field_end):
        field_run = OxmlElement("w:r")
        field_run.append(field_node)
        right._p.append(field_run)


def add_cover(doc: Document) -> None:
    for _ in range(4):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("ASSIGNMENT 1 | FINAL IMPLEMENTATION GUIDE")
    set_run_font(run, size=10.5, color=GREEN, bold=True)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(9)
    run = title.add_run("JalDrishti")
    set_run_font(run, size=32, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("AI-based Village Pond Planning System")
    set_run_font(run, size=18, color=DEEP_GREEN, bold=True)
    strapline = doc.add_paragraph()
    strapline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    strapline.paragraph_format.space_after = Pt(30)
    run = strapline.add_run("Complete design, implementation, algorithms, testing, deployment, and viva reference")
    set_run_font(run, size=11.5, color=GRAY, italic=True)
    add_callout(doc, "DOCUMENT PURPOSE", "A beginner-friendly but technically complete explanation of how the application turns contour or live geospatial evidence into ranked pond-screening options. It documents the deployed release and its limits; it is not a construction drawing or excavation approval.", fill=MINT)
    metadata = [
        ("Student", "Sneha Nagmoti"),
        ("Course deliverable", "Assignment 1 - Phase 3 final implementation"),
        ("Verified release", "0a385c0"),
        ("Verification date", "29 August 2026"),
        ("Frontend", "sneha-village-pond-planning-2026.onrender.com"),
        ("Repository", "github.com/snehanagmoti/Village-Pond-planning"),
    ]
    add_table(doc, ["Document field", "Value"], metadata, [2700, 6660], font_size=9.2)
    doc.add_page_break()


def add_document_control(doc: Document) -> None:
    add_heading(doc, "Document control and how to use this guide", 1, "sec_document_control", 1)
    add_rich_text(doc.add_paragraph(), "This guide is written so that a reader with no GIS or hydrology background can follow the system from input to output. Read Sections 1-4 for the big picture, Sections 5-10 for the algorithms, Sections 11-15 for engineering and software details, and the appendices for demonstration and viva preparation.")
    add_callout(doc, "ACADEMIC INTEGRITY", "AI-assisted development was used for review, implementation support, testing, and document formatting. The submitted algorithms and measurements were inspected and exercised with automated and production tests. The student remains responsible for understanding and explaining every component.", fill=PALE_GOLD, border="8A6400")
    add_table(
        doc,
        ["Item", "Verified state"],
        [
            ("Scope", "Code, algorithms, UI, tests, deployment, and documentation"),
            ("Functional release", "0a385c0 deployed on both Render services"),
            ("Backend verification", "68 tests passed; 88.79% statement coverage; Ruff passed"),
            ("Frontend verification", "11 tests passed; Oxlint passed; Vite production build passed"),
            ("Production contour verification", "KML, KMZ, automatic, point, region, and invalid-point cases exercised"),
            ("Design status", "Screening prototype; field and qualified engineering verification required"),
        ],
        [2700, 6660],
    )
    doc.add_page_break()


def parse_report_sections(lines: list[str]) -> list[tuple[int, str, str]]:
    headings: list[tuple[int, str, str]] = []
    start = False
    count = 10
    for line in lines:
        if line.startswith("## 1. Abstract"):
            start = True
        if not start:
            continue
        match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if match:
            markdown_level = len(match.group(1))
            level = markdown_level - 1
            count += 1
            headings.append((level, clean_md(match.group(2)), f"sec_{count:03d}"))
    return headings


def add_toc(doc: Document, headings: list[tuple[int, str, str]]) -> None:
    add_heading(doc, "Table of contents", 1, "toc", 2)
    intro = doc.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    add_rich_text(intro, "The entries below are internal links in the editable Word document.", size=9.5, color=GRAY, italic=True)
    for level, title, bookmark in headings:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.0 if level == 1 else 0.25 if level == 2 else 0.5)
        paragraph.paragraph_format.space_after = Pt(3)
        add_internal_link(paragraph, bookmark, title)


def markdown_table(lines: list[str], start: int) -> tuple[int, list[str], list[list[str]]]:
    headers = [cell.strip() for cell in lines[start].strip().strip("|").split("|")]
    rows: list[list[str]] = []
    index = start + 2
    while index < len(lines) and lines[index].strip().startswith("|"):
        rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
        index += 1
    return index, headers, rows


def table_widths(columns: int) -> list[int]:
    patterns = {
        2: [3300, 6060],
        3: [1900, 3730, 3730],
        4: [1550, 3110, 2350, 2350],
        5: [1200, 2240, 2000, 1960, 1960],
        6: [1560, 1560, 1560, 1560, 1560, 1560],
    }
    return patterns.get(columns, [TABLE_WIDTH_DXA // columns] * (columns - 1) + [TABLE_WIDTH_DXA - (TABLE_WIDTH_DXA // columns) * (columns - 1)])


def add_report_body(doc: Document, lines: list[str], headings: list[tuple[int, str, str]], bullet_id: int, number_id: int, decimal_abstract_id: int) -> None:
    heading_iter = iter(headings)
    next_heading = next(heading_iter, None)
    started = False
    paragraph_lines: list[str] = []
    figure_number = 1
    bookmark_id = 20
    figure_map = {
        "3. System architecture": [("architecture.png", "System architecture and service boundaries", "Block diagram showing React and Leaflet, FastAPI, contour and live-source services, the shared hydrology core, optional database, and external providers.")],
        "4. Phase 2 contour-file analysis": [("workflow_comparison.png", "How contour upload and live analysis complement each other", "Comparison of the contour-file and live-location workflows feeding a shared output contract.")],
        "4.3 Hydrology and candidate selection": [("contour_pipeline.png", "End-to-end contour upload algorithm", "Eight-step flowchart from safe parsing to selected catchment, rainfall, runoff, and pond geometry."), ("safety_gates.png", "Hard candidate eligibility gates", "Decision flow showing study boundary, setback, outlet, water buffer, and upstream-catchment checks before ranking.")],
        "4.4 Provided-map result": [("production_options.png", "Three ranked options from the supplied KML", "Three grouped bar panels compare suitability, upstream catchment area, and detected-water clearance for the production alternatives."), ("production-map-expanded.png", "Expanded production map with panel collapsed", "Satellite map displaying uploaded contour extent, reconstructed contours, selected catchment, drainage path, hydrology outlet, and three numbered pond alternatives.")],
        "6. Rainfall analysis": [("rainfall_climatology.png", "Historical monthly rainfall at the contour study", "Bar chart of monthly climatological rainfall with annual mean and valid-year count." )],
        "9. Frontend and visualization": [("production-home.png", "Final home and live-analysis interface", "Full application view with India satellite map and the responsive course-project analysis panel."), ("production-live-result.png", "Final live-analysis result workspace", "Live analysis result showing the mapped catchment and pond options beside the concise completion summary, collapsed technical notes, and source-quality labels." )],
    }

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            paragraph = doc.add_paragraph()
            add_rich_text(paragraph, " ".join(part.strip() for part in paragraph_lines))
            paragraph_lines = []

    index = 0
    in_code = False
    code_lines: list[str] = []
    active_number_id = number_id
    last_was_number = False
    while index < len(lines):
        line = lines[index]
        if line.startswith("## 1. Abstract"):
            started = True
        if not started:
            index += 1
            continue
        if line.startswith("```"):
            flush_paragraph()
            last_was_number = False
            if in_code:
                add_code(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", line)
        if heading_match:
            flush_paragraph()
            last_was_number = False
            level = len(heading_match.group(1)) - 1
            title = clean_md(heading_match.group(2))
            if next_heading is None:
                raise RuntimeError("Heading index exhausted")
            _, _, bookmark = next_heading
            add_heading(doc, title, level, bookmark, bookmark_id, page_break=(level == 1))
            bookmark_id += 1
            next_heading = next(heading_iter, None)
            for asset_name, caption, alt in figure_map.get(title, []):
                add_figure(doc, ASSETS / asset_name, caption, alt, figure_number)
                figure_number += 1
            index += 1
            continue
        if line.strip().startswith("|") and index + 1 < len(lines) and re.match(r"^\s*\|?\s*:?-+", lines[index + 1]):
            flush_paragraph()
            last_was_number = False
            index, headers, rows = markdown_table(lines, index)
            add_table(doc, headers, rows, table_widths(len(headers)))
            continue
        if re.match(r"^\s*-\s+", line):
            flush_paragraph()
            last_was_number = False
            item_text = re.sub(r"^\s*-\s+", "", line)
            next_index = index + 1
            while next_index < len(lines) and re.match(r"^\s{2,}\S", lines[next_index]) and not re.match(r"^\s*[-\d]+\.\s+|^\s*-\s+", lines[next_index]):
                item_text += " " + lines[next_index].strip()
                next_index += 1
            add_bullet(doc, item_text, bullet_id)
            index = next_index
            continue
        if re.match(r"^\s*\d+\.\s+", line):
            flush_paragraph()
            if not last_was_number:
                active_number_id = clone_numbering(doc, decimal_abstract_id)
            item_text = re.sub(r"^\s*\d+\.\s+", "", line)
            next_index = index + 1
            while next_index < len(lines) and re.match(r"^\s{2,}\S", lines[next_index]) and not re.match(r"^\s*\d+\.\s+|^\s*-\s+", lines[next_index]):
                item_text += " " + lines[next_index].strip()
                next_index += 1
            add_number(doc, item_text, active_number_id)
            last_was_number = True
            index = next_index
            continue
        if not line.strip():
            flush_paragraph()
            last_was_number = False
        else:
            paragraph_lines.append(line)
            last_was_number = False
        index += 1
    flush_paragraph()


def add_appendices(doc: Document, bullet_id: int, number_id: int, decimal_abstract_id: int, production: dict) -> None:
    bookmark = 200
    add_heading(doc, "Appendix A. Algorithms explained from first principles", 1, "appendix_algorithms", bookmark, page_break=True)
    bookmark += 1
    add_heading(doc, "A.1 What a contour line means", 2, "appendix_contour_meaning", bookmark)
    bookmark += 1
    add_rich_text(doc.add_paragraph(), "A contour line connects locations having the same elevation. Lines close together usually mean a steeper surface; lines farther apart usually mean a gentler surface. A KML contour map is still only a set of lines. Water-flow algorithms need an elevation value for every grid cell, so the system first reconstructs the unknown cells between lines.")
    add_callout(doc, "KML VERSUS KMZ", "KML is readable XML containing placemarks and coordinates. KMZ is a ZIP archive whose main content is KML, commonly named doc.kml. Compressing KML to KMZ changes storage and transfer size; it does not improve terrain accuracy. The deployed API accepts and validates both formats.")
    add_heading(doc, "A.2 Harmonic interpolation", 2, "appendix_interpolation", bookmark)
    bookmark += 1
    add_rich_text(doc.add_paragraph(), "The rasterized contour cells are fixed observations. Every unknown interior cell is initialized from nearby observations and repeatedly replaced by the mean of valid neighboring cells. The fixed cells never move. Repetition creates a smooth surface that exactly preserves the rasterized contours. The process stops when the maximum update is below tolerance or the iteration limit is reached.")
    add_bullet(doc, "Why use it: deterministic, bounded, easy to audit, and appropriate for smoothly varying terrain between dense contours.", bullet_id)
    add_bullet(doc, "What it cannot know: breaks of slope, channels, embankments, and survey errors not represented by the contour geometry.", bullet_id)
    add_heading(doc, "A.3 Priority-Flood depression conditioning", 2, "appendix_priority_flood", bookmark)
    bookmark += 1
    add_rich_text(doc.add_paragraph(), "Digital surfaces often contain one-cell pits caused by sampling or interpolation. A naive D8 router would trap water in those cells. Priority-Flood begins at the analysis boundary and visits cells from low to high elevation using a priority queue. If an unvisited neighbor is lower than the current spill elevation, it is raised only enough to drain. The output is a connected drainage surface, while the report records how much of the grid changed so large corrections remain visible.")
    add_heading(doc, "A.4 D8 flow direction, accumulation, and reverse catchment", 2, "appendix_d8", bookmark)
    bookmark += 1
    add_rich_text(doc.add_paragraph(), "D8 means each cell may drain to one of its eight neighbors. The implementation chooses the lower neighbor with the greatest elevation drop divided by travel distance. Cardinal neighbors are one cell away and diagonal neighbors are √2 cells away. A tiny deterministic gradient resolves perfectly equal flats without random results.")
    algorithm_number_id = clone_numbering(doc, decimal_abstract_id)
    add_number(doc, "Flow direction creates a directed graph: each cell has at most one downstream neighbor.", algorithm_number_id)
    add_number(doc, "Topological accumulation starts every valid cell with one contributing cell and passes that total downstream.", algorithm_number_id)
    add_number(doc, "After a pond point is selected, reverse traversal follows upstream links and collects only cells that drain to that point.", algorithm_number_id)
    add_number(doc, "Cell area is corrected for latitude and multiplied by collected-cell count to obtain square metres and hectares.", algorithm_number_id)
    add_callout(doc, "THE CRITICAL CORRECTION", "The reported catchment is delineated upstream of the selected pond option, not at the map-edge hydrology outlet. Selecting a different option changes the catchment and therefore changes runoff and pond geometry.", fill=PALE_GOLD, border="8A6400")
    add_heading(doc, "A.5 Multi-criteria pond ranking", 2, "appendix_ranking", bookmark)
    bookmark += 1
    add_rich_text(doc.add_paragraph(), "No single terrain variable is sufficient. A large catchment may be on a steep or unsafe cell; a flat cell may collect almost no water. Therefore the system first applies hard exclusions and then combines normalized evidence. Logarithmic catchment area is used because raw accumulation can span orders of magnitude and would otherwise dominate every other factor.")
    add_table(
        doc,
        ["Criterion", "With water evidence", "Without water evidence", "Reason"],
        [
            ("Log upstream area", "52%", "58%", "Favors useful contributing area without letting very large basins overwhelm the score"),
            ("Local flatness", "20%", "22%", "Favors gentler local terrain for preliminary siting"),
            ("Lower relative elevation", "10%", "11%", "Prefers natural receiving locations within the study area"),
            ("Boundary clearance", "8%", "9%", "Reduces truncated or edge-dominated selections"),
            ("Detected-water clearance", "10%", "Not available", "Rewards separation after the hard 60 m exclusion"),
        ],
        [2100, 1450, 1600, 4210],
    )
    add_rich_text(doc.add_paragraph(), "After sorting by score, non-maximum suppression keeps alternatives at least 100 m or three cells apart. This prevents three nearly identical markers from being presented as meaningful choices.")

    add_heading(doc, "Appendix B. River and water edge cases", 1, "appendix_river", bookmark, page_break=True)
    bookmark += 1
    add_rich_text(doc.add_paragraph(), "A pond should not be recommended directly in a river channel. The implementation therefore aligns the satellite water mask to the terrain grid, expands it by a metric buffer, and removes those cells before ranking. Manual point and region modes use the same eligibility mask; they cannot bypass the rule.")
    add_table(
        doc,
        ["Edge case", "Implemented response", "Remaining real-world check"],
        [
            ("Wide visible river", "Water pixels and 60 m surroundings are excluded", "Confirm bankfull width, floodplain, and statutory setback"),
            ("Narrow or shaded stream", "Warning states that non-detection is not absence", "Field walk, hydrography data, and local drainage knowledge"),
            ("Muddy or seasonal channel", "May be missed by RGB/HSV colour classification", "Multi-season imagery and monsoon evidence"),
            ("Canal, culvert, road drain", "May alter real flow but not appear in the contour DEM", "Survey crossings and drainage structures"),
            ("Point placed on detected water", "HTTP 422 invalid_contour_selection", "Do not override without qualified investigation"),
            ("Region contains water and land", "Only eligible land cells inside the polygon are ranked", "Verify the selected parcel and access"),
        ],
        [2100, 3150, 4110],
    )
    add_callout(doc, "WHY THE RESULT IS STILL DEGRADED", "A conservative colour mask reduces obvious false positives but cannot establish a legal river boundary or construction clearance. The degraded status is an honest quality label, not a software failure.", fill=PALE_RED, border="A13B3B")

    add_heading(doc, "Appendix C. Complete output dictionary", 1, "appendix_outputs", bookmark, page_break=True)
    bookmark += 1
    output_rows = [
        ("analysis_status", "reliable / degraded / incomplete", "Overall evidence status; contour interpolation remains degraded by design"),
        ("contour_summary", "counts, range, interval", "What the KML/KMZ contained and whether it had sufficient terrain evidence"),
        ("grid", "rows, columns, cell size, convergence", "Quality and resolution of the reconstructed surface"),
        ("pond_location", "coordinate, elevation, slope, score", "The selected eligible grid cell and how it was chosen"),
        ("candidate_options", "ranked array", "Spatially separated alternatives with upstream area and water clearance"),
        ("selection", "automatic / point / region", "User request, snapped distance, and optional region vertices"),
        ("outlet_location", "coordinate", "Hydrology evidence at the map edge; explicitly not a pond recommendation"),
        ("catchment", "boundary, m², ha, cells", "Every selected grid cell that drains to the pond point"),
        ("rainfall_data", "annual mean + 12 months", "Complete-year climatology and valid-year counts"),
        ("runoff_stats", "C, A, P, annual volume", "Screening water yield V = C × A × P"),
        ("pond", "depth, dimensions, capacity", "Preliminary water and excavation geometry, including freeboard"),
        ("water_screening", "ratio, buffer, method", "How detected water affected eligibility"),
        ("contours / drainage_path", "map arrays", "Evidence layers displayed by Leaflet"),
        ("quality.sources", "provenance records", "Provider, retrieval date, resolution, model, coverage, license link"),
        ("quality.warnings", "plain-language array", "Known omissions and required verification"),
    ]
    add_table(doc, ["Field", "Contents", "Meaning"], output_rows, [2300, 2500, 4560], font_size=8.5)
    add_heading(doc, "C.1 Exact automatic production output", 2, "appendix_output_example", bookmark)
    bookmark += 1
    selected = next(item for item in production["candidate_options"] if item["selected"])
    add_table(
        doc,
        ["Output", "Production value"],
        [
            ("Selected point", f"{selected['lat']:.6f}, {selected['lng']:.6f}"),
            ("Suitability", f"{selected['suitability_score']:.2f} / 100"),
            ("Selected catchment", f"{production['catchment']['area_hectares']:.4f} ha"),
            ("Historical rainfall", f"{production['rainfall_data']['annual_avg_mm']:.2f} mm/year; {production['rainfall_data']['valid_years']} years"),
            ("Annual runoff", f"{production['runoff_stats']['estimated_volume_m3']:,.2f} m³/year"),
            ("Pond capacity", f"{production['pond']['capacity_m3']:,.2f} m³"),
            ("Excavation volume", f"{production['pond']['excavation_volume_m3']:,.2f} m³"),
            ("Water screening", f"{production['water_screening']['detected_water_ratio'] * 100:.2f}% detected; {production['water_screening']['exclusion_buffer_m']:.0f} m buffer"),
        ],
        [3300, 6060],
    )

    add_heading(doc, "Appendix D. API examples and error behavior", 1, "appendix_api", bookmark, page_break=True)
    bookmark += 1
    add_heading(doc, "D.1 Automatic contour analysis", 2, "appendix_api_auto", bookmark)
    bookmark += 1
    add_code(doc, "curl.exe -X POST ^\n  -F \"contour_file=@C:\\path\\contours_1m.kml\" ^\n  -F \"selection_mode=automatic\" ^\n  https://sneha-village-pond-api-2026.onrender.com/api/analyze-contour")
    add_heading(doc, "D.2 Manual point", 2, "appendix_api_point", bookmark)
    bookmark += 1
    add_code(doc, "curl.exe -X POST ^\n  -F \"contour_file=@C:\\path\\contours_1m.kml\" ^\n  -F \"selection_mode=point\" ^\n  -F \"selected_lat=21.245156\" ^\n  -F \"selected_lng=81.289215\" ^\n  https://sneha-village-pond-api-2026.onrender.com/api/analyze-contour")
    add_heading(doc, "D.3 Region-constrained search", 2, "appendix_api_region", bookmark)
    bookmark += 1
    add_code(doc, "selected_region=[\n  {\"lat\":21.2424,\"lng\":81.2864},\n  {\"lat\":21.2434,\"lng\":81.2864},\n  {\"lat\":21.2434,\"lng\":81.2875},\n  {\"lat\":21.2424,\"lng\":81.2875}\n]")
    add_table(
        doc,
        ["HTTP", "Stable code", "Meaning"],
        [
            ("200", "success", "Typed screening response returned"),
            ("413", "contour_file_too_large", "Compressed upload exceeds configured maximum"),
            ("422", "invalid_contour_file", "Unsafe, malformed, or insufficient contour content"),
            ("422", "invalid_contour_selection", "Point/region is malformed or fails an eligibility gate"),
            ("429", "rate_limit_exceeded", "Client request budget exhausted"),
            ("500", "contour_analysis_failed", "Unexpected server-side processing error with a safe message"),
        ],
        [1000, 2800, 5560],
    )

    add_heading(doc, "Appendix E. Manual test script", 1, "appendix_manual_test", bookmark, page_break=True)
    bookmark += 1
    test_groups = [
        ("Application shell", ["Open the production URL and confirm satellite imagery, live/contour tabs, decision-support warning, keyboard focus, and the Hide panel control.", "Collapse and reopen the panel; verify that the map expands and all layers remain aligned.", "Resize to a narrow window; verify that controls remain readable and no content becomes unreachable."]),
        ("Automatic KML", ["Upload contours_1m.kml and confirm the filename, progress message, and successful contour screening report.", "Confirm all 1,355 contours are represented, the uploaded extent is visible, and the legend distinguishes catchment, contours, drainage path, outlet, and options.", "Confirm three option cards, 352.9523 ha selected catchment, 1,280.13 mm rainfall, 1,355,474.66 m³ runoff, and 1,084,379.73 m³ capacity."]),
        ("Alternative point", ["Choose option 3 and recompute.", "Confirm the selected marker changes and catchment becomes 301.1143 ha.", "Confirm runoff becomes 1,156,396.32 m³ and capacity becomes 925,117.06 m³."]),
        ("Region selection", ["Activate Draw region and create a polygon around option 2.", "Confirm the chosen point is inside the polygon and the catchment becomes 381.4633 ha.", "Confirm clearing the selection returns the workflow to automatic mode."]),
        ("Negative and safety", ["Attempt a point outside the uploaded extent; verify a clear rejection rather than a result.", "Attempt a point inside a detected-water buffer in a controlled test; verify invalid_contour_selection.", "Read the warning that satellite non-detection does not prove a river is absent."]),
        ("KMZ and live analysis", ["Upload the KMZ equivalent and confirm the same automatic output.", "Run live analysis at 21.24, 81.29 with a 2 km radius.", "Confirm three live options, 35 rainfall years, annual runoff, pond geometry, sources, and degraded-quality explanations."]),
    ]
    for group, tests in test_groups:
        add_heading(doc, group, 2, f"test_{bookmark}", bookmark)
        bookmark += 1
        for test in tests:
            add_bullet(doc, "□ " + test, bullet_id)

    add_heading(doc, "Appendix F. Viva preparation", 1, "appendix_viva", bookmark, page_break=True)
    bookmark += 1
    viva_rows = [
        ("Why is contour output degraded?", "Because the raster is interpolated between vector isolines and is not equivalent to a surveyed DEM."),
        ("Why Priority-Flood?", "It removes artificial sinks in a deterministic, efficient way so the D8 network can drain."),
        ("What is D8?", "A single-flow-direction model where each cell drains to its steepest lower neighbor among eight choices."),
        ("What was the major catchment fix?", "Reverse traversal now starts at the selected pond cell, not the map-edge outlet."),
        ("Why offer three options?", "A comparative model should expose spatial alternatives; one maximum is not an engineering approval."),
        ("Can a user force any point?", "No. Clicked points are snapped and still must pass boundary, outlet, water, and catchment gates."),
        ("How are rivers handled?", "Detected water plus 60 m is excluded, but field/hydrography verification remains mandatory because RGB can miss channels."),
        ("How is annual runoff calculated?", "V = C × A × P, with P converted from millimetres to metres."),
        ("Is rainfall a design storm?", "No. It is a multi-year climatological mean; peak-flow and spillway design need approved intensity-duration-frequency data."),
        ("What does pond capacity mean?", "Preliminary storage to water level from rectangular-frustum geometry, constrained by available area."),
        ("Why not claim government land?", "Satellite colour cannot prove ownership, tenure, consent, or legal availability."),
        ("How was the result verified?", "68 backend tests, 11 frontend tests, lint/build gates, exact KML/KMZ production uploads, and manual selection checks."),
    ]
    add_table(doc, ["Question", "Short answer"], viva_rows, [3300, 6060], font_size=8.7)

    add_heading(doc, "Appendix G. Glossary", 1, "appendix_glossary", bookmark, page_break=True)
    glossary = [
        ("Catchment / watershed", "Land area whose modelled drainage reaches the selected point"),
        ("Contour", "Line connecting places with equal elevation"),
        ("DEM", "Digital elevation model: a grid of ground elevations"),
        ("D8", "Eight-neighbor single-flow-direction routing method"),
        ("Flow accumulation", "Count or area of upstream cells contributing through a cell"),
        ("Freeboard", "Vertical margin between design water level and crest"),
        ("Harmonic interpolation", "Smoothly estimates unknown grid cells while preserving fixed observations"),
        ("Hydrology outlet", "Terminal drainage cell at the analysis boundary; not automatically a pond site"),
        ("KML", "XML geospatial format standardized by OGC"),
        ("KMZ", "Compressed ZIP package containing KML"),
        ("Priority-Flood", "Priority-queue algorithm for conditioning depressions in a DEM"),
        ("Runoff coefficient C", "Fraction of rainfall volume assumed to become direct runoff"),
        ("Screening", "Preliminary comparative decision support, not an approved final design"),
        ("Suitability score", "Normalized comparative score after all hard eligibility checks"),
    ]
    add_table(doc, ["Term", "Plain-language meaning"], glossary, [2700, 6660], font_size=9)

    add_heading(doc, "Appendix H. Submission links and final checklist", 1, "appendix_submission", bookmark, page_break=True)
    add_callout(doc, "LIVE FRONTEND", "https://sneha-village-pond-planning-2026.onrender.com", fill=MINT)
    add_callout(doc, "GITHUB REPOSITORY", "https://github.com/snehanagmoti/Village-Pond-planning", fill=PALE_BLUE, border=BLUE)
    add_callout(doc, "INTERACTIVE API DOCUMENTATION", "https://sneha-village-pond-api-2026.onrender.com/docs", fill=PALE_GOLD, border="8A6400")
    for item in [
        "Source code is committed and pushed.",
        "Both Render services report the verified release deployed successfully.",
        "The production frontend opens and both workflows return complete results.",
        "The provided KML and its KMZ equivalent return HTTP 200.",
        "Automatic, point, region, and invalid-selection behavior was checked.",
        "The editable documentation and final PDF were rendered and visually inspected.",
        "The student has reviewed the limitations and can explain the algorithms in the viva.",
    ]:
        add_bullet(doc, "□ " + item, bullet_id)


def main() -> None:
    lines = REPORT.read_text(encoding="utf-8").splitlines()
    headings = parse_report_sections(lines)
    appendix_toc = [
        (1, "Appendix A. Algorithms explained from first principles", "appendix_algorithms"),
        (1, "Appendix B. River and water edge cases", "appendix_river"),
        (1, "Appendix C. Complete output dictionary", "appendix_outputs"),
        (1, "Appendix D. API examples and error behavior", "appendix_api"),
        (1, "Appendix E. Manual test script", "appendix_manual_test"),
        (1, "Appendix F. Viva preparation", "appendix_viva"),
        (1, "Appendix G. Glossary", "appendix_glossary"),
        (1, "Appendix H. Submission links and final checklist", "appendix_submission"),
    ]
    with PRODUCTION.open(encoding="utf-8") as handle:
        production = json.load(handle)
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_id, number_id, _bullet_abstract_id, decimal_abstract_id = add_custom_numbering(doc)
    doc.core_properties.title = "JalDrishti - Complete Project Documentation"
    doc.core_properties.subject = "AI-based Village Pond Planning System"
    doc.core_properties.keywords = "pond planning, KML, KMZ, catchment, D8, rainfall, runoff"
    add_cover(doc)
    add_document_control(doc)
    add_toc(doc, headings + appendix_toc)
    add_report_body(doc, lines, headings, bullet_id, number_id, decimal_abstract_id)
    add_appendices(doc, bullet_id, number_id, decimal_abstract_id, production)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
